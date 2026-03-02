"""
src/ingestion/extract_text.py
PDF / DOCX / TXT text extraction with OCR fallback and text cleaning.
"""
import logging
import os
from typing import List

from src.ingestion.text_cleaner import clean_text

logger = logging.getLogger(__name__)


class TextExtractor:
    def __init__(self):
        pass

    # ── Primary PDF extractor ─────────────────────────────────────────────────
    def extract_from_pdf_advanced(self, pdf_path: str) -> List[dict]:
        """
        Extract text from each page of a PDF.
        Tries PyMuPDF (fitz) first; falls back to pdfplumber for better table support,
        then finally OCR for image-only PDFs.
        """
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF not found: {pdf_path}")

        # 1. Try PyMuPDF (Fast, preserves some layout)
        docs = self._extract_with_pymupdf(pdf_path)

        # 2. Try pdfplumber if PyMuPDF yields too little text
        if not docs or len("".join([d["page_content"] for d in docs])) < 100:
            logger.info(f"PyMuPDF extracted minimal text from {os.path.basename(pdf_path)}; trying pdfplumber.")
            plumber_docs = self._extract_with_pdfplumber(pdf_path)
            if plumber_docs:
                docs = plumber_docs

        if docs:
            return docs

        # 3. Last resort: OCR
        logger.info(f"No text extracted by standard methods; attempting OCR on {os.path.basename(pdf_path)}.")
        return self._ocr_fallback(pdf_path)

    def extract_from_docx(self, docx_path: str) -> List[dict]:
        """
        Extract text from DOCX files — paragraphs AND table cells.
        Preserves table content that was previously silently ignored.
        """
        try:
            import docx
            doc = docx.Document(docx_path)
            full_text: List[str] = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Bug 13 Fix: Also extract table cell content (was ignored before)
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        full_text.append(" | ".join(row_texts))

            text = "\n".join(full_text)
            text = clean_text(text)
            if not text:
                logger.warning(f"DOCX extraction yielded empty text: {docx_path}")
                return []

            return [{
                "page_content": text,
                "metadata": {
                    "type": "docx",
                    "source": os.path.basename(docx_path),
                    "page": 1,
                }
            }]
        except Exception as e:
            logger.error(f"DOCX extraction error for {docx_path}: {e}")
            return []

    # ── PyMuPDF ──────────────────────────────────────────────────────────────
    def _extract_with_pymupdf(self, pdf_path: str) -> List[dict]:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF (fitz) not installed — skipping.")
            return []
        try:
            doc = fitz.open(pdf_path)
            docs = []
            for i, page in enumerate(doc):
                raw = page.get_text("text")  # "text" mode preserves layout better
                text = clean_text(raw)
                if text:
                    docs.append({
                        "page_content": text,
                        "metadata": {
                            "type": "text",
                            "source": os.path.basename(pdf_path),
                            "page": i + 1,
                        },
                    })
            doc.close()
            return docs
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed for {pdf_path}: {e}")
            return []

    # ── pdfplumber fallback ───────────────────────────────────────────────────
    def _extract_with_pdfplumber(self, pdf_path: str) -> List[dict]:
        """Better at tables and multi-column layouts."""
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber not installed — skipping.")
            return []
        try:
            docs = []
            with pdfplumber.open(pdf_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    raw = page.extract_text() or ""
                    text = clean_text(raw)
                    if text:
                        docs.append({
                            "page_content": text,
                            "metadata": {
                                "type": "text_pdfplumber",
                                "source": os.path.basename(pdf_path),
                                "page": i + 1,
                            },
                        })
            return docs
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed for {pdf_path}: {e}")
            return []

    # ── OCR fallback ──────────────────────────────────────────────────────────
    def _ocr_fallback(self, pdf_path: str) -> List[dict]:
        """Attempt OCR on each page if no text was extracted (scanned PDFs)."""
        try:
            import fitz
            from PIL import Image
            import pytesseract
        except ImportError:
            logger.warning("OCR dependencies (fitz/PIL/pytesseract) not available — skipping OCR.")
            return []

        # Set tesseract binary on Windows if not on PATH
        try:
            import shutil as _sh
            if _sh.which("tesseract") is None:
                candidate = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                if os.path.exists(candidate):
                    pytesseract.pytesseract.tesseract_cmd = candidate
        except Exception:
            pass

        docs: List[dict] = []
        try:
            doc = fitz.open(pdf_path)
            for i, page in enumerate(doc):
                pix = page.get_pixmap(alpha=False, dpi=200)  # higher DPI = better OCR
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                raw = pytesseract.image_to_string(img, config="--psm 6")
                text = clean_text(raw)
                if text:
                    docs.append({
                        "page_content": text,
                        "metadata": {
                            "type": "ocr",
                            "source": os.path.basename(pdf_path),
                            "page": i + 1,
                        },
                    })
            doc.close()
        except Exception as e:
            logger.error(f"OCR fallback failed for {pdf_path}: {e}")
            return []
        return docs

    # ── Simple single-string extractor (for non-PDF, legacy) ─────────────────
    def extract_from_pdf_simple(self, pdf_path: str) -> str:
        try:
            import fitz
            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return clean_text(text)
        except Exception as e:
            logger.error(f"extract_from_pdf_simple failed: {e}")
            return ""
